"""
Insert real DOCX footnotes into an existing document.
"""
from __future__ import annotations

import copy
import os
import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml.ns import qn
from lxml import etree


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
FOOTNOTE_MARKER_PATTERN = re.compile(r"\[\[(FN_[^\]]+)\]\]")

MINIMAL_FOOTNOTES_XML = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:footnotes xmlns:w="{W}">
  <w:footnote w:type="separator" w:id="2">
    <w:p>
      <w:r>
        <w:separator/>
      </w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="3">
    <w:p>
      <w:r>
        <w:continuationSeparator/>
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>
"""


class FootnoteInserter:
    """Low-level DOCX footnote writer."""

    def __init__(self, docx_path: str):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")

        self.doc = Document(docx_path)
        self.paragraphs = [para for para in self.doc.paragraphs if para.text.strip()]
        self.ref_style_id, self.text_style_id = self._detect_footnote_styles()

    def _detect_footnote_styles(self) -> Tuple[Optional[str], Optional[str]]:
        """Try to reuse built-in footnote styles when present."""
        ref_style_id = None
        text_style_id = None

        try:
            styles_part = self.doc.part.styles
            for style_elem in styles_part._element.iter(qn("w:style")):
                name_elem = style_elem.find(qn("w:name"))
                if name_elem is None:
                    continue
                name_val = name_elem.get(qn("w:val"), "").lower()
                style_id = style_elem.get(qn("w:styleId"))
                if name_val == "footnote reference":
                    ref_style_id = style_id
                elif name_val == "footnote text":
                    text_style_id = style_id
        except Exception:
            pass

        return ref_style_id, text_style_id

    def _get_or_create_footnotes_tree(self):
        """Return the footnotes tree and part, creating them when absent."""
        for rel in self.doc.part.rels.values():
            if rel.reltype == RT.FOOTNOTES:
                fn_part = rel.target_part
                fn_tree = etree.fromstring(fn_part.blob)
                return fn_tree, fn_part

        fn_tree = etree.fromstring(MINIMAL_FOOTNOTES_XML.encode("utf-8"))
        fn_part = XmlPart(
            PackURI("/word/footnotes.xml"),
            CT.WML_FOOTNOTES,
            fn_tree,
            self.doc.part.package,
        )
        self.doc.part.relate_to(fn_part, RT.FOOTNOTES)
        self._ensure_settings_footnote_pr()
        return fn_tree, fn_part

    def _ensure_settings_footnote_pr(self) -> None:
        """Ensure settings.xml references the separator footnotes."""
        settings_part = self.doc.part._settings_part
        settings_elem = settings_part.element
        footnote_pr = settings_elem.find(qn("w:footnotePr"))

        if footnote_pr is None:
            footnote_pr = etree.SubElement(settings_elem, qn("w:footnotePr"))

        existing_ids = {
            child.get(qn("w:id"))
            for child in footnote_pr.findall(qn("w:footnote"))
        }

        for footnote_id in ("2", "3"):
            if footnote_id in existing_ids:
                continue
            child = etree.SubElement(footnote_pr, qn("w:footnote"))
            child.set(qn("w:id"), footnote_id)

    def _get_next_footnote_id(self, fn_tree) -> int:
        """Return the next available non-reserved footnote ID."""
        existing_ids = set()
        for footnote in fn_tree:
            footnote_type = footnote.get(f"{{{W}}}type")
            footnote_id = footnote.get(f"{{{W}}}id")
            if footnote_type is None and footnote_id is not None:
                try:
                    existing_ids.add(int(footnote_id))
                except ValueError:
                    pass

        next_id = max(existing_ids, default=0) + 1
        while next_id in (2, 3):
            next_id += 1
        return next_id

    def _add_footnote_to_tree(self, fn_tree, docx_id: int, footnote_text: str) -> None:
        """Append a new footnote node to footnotes.xml."""
        footnote = etree.SubElement(fn_tree, f"{{{W}}}footnote")
        footnote.set(f"{{{W}}}id", str(docx_id))

        paragraph = etree.SubElement(footnote, f"{{{W}}}p")

        p_pr = etree.SubElement(paragraph, f"{{{W}}}pPr")
        p_style = etree.SubElement(p_pr, f"{{{W}}}pStyle")
        p_style.set(f"{{{W}}}val", self.text_style_id or "FootnoteText")

        r_num = etree.SubElement(paragraph, f"{{{W}}}r")
        r_pr_num = etree.SubElement(r_num, f"{{{W}}}rPr")
        if self.ref_style_id:
            r_style_num = etree.SubElement(r_pr_num, f"{{{W}}}rStyle")
            r_style_num.set(f"{{{W}}}val", self.ref_style_id)
        else:
            vert_align = etree.SubElement(r_pr_num, f"{{{W}}}vertAlign")
            vert_align.set(f"{{{W}}}val", "superscript")
        etree.SubElement(r_num, f"{{{W}}}footnoteRef")

        r_space = etree.SubElement(paragraph, f"{{{W}}}r")
        t_space = etree.SubElement(r_space, f"{{{W}}}t")
        t_space.set(XML_SPACE, "preserve")
        t_space.text = " "

        r_text = etree.SubElement(paragraph, f"{{{W}}}r")
        t_text = etree.SubElement(r_text, f"{{{W}}}t")
        t_text.set(XML_SPACE, "preserve")
        t_text.text = footnote_text

    def _save_footnotes_tree(self, fn_tree, fn_part) -> None:
        """Persist the modified footnotes tree back into the package."""
        fn_part._blob = etree.tostring(
            fn_tree,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

    def _insert_footnote_at_location(self, para, char_pos: int, docx_fn_id: int) -> None:
        """Insert a footnote reference at the character position within a paragraph."""
        p_elem = para._element
        cumulative = 0
        target_run_elem = None
        offset_in_run = 0
        run_position_in_parent = 0

        children = list(p_elem)
        for idx, child in enumerate(children):
            if child.tag != f"{{{W}}}r":
                continue

            t_elem = child.find(f"{{{W}}}t")
            run_text = t_elem.text if (t_elem is not None and t_elem.text) else ""
            run_len = len(run_text)

            if cumulative + run_len >= char_pos:
                target_run_elem = child
                offset_in_run = char_pos - cumulative
                run_position_in_parent = idx
                break
            cumulative += run_len

        if target_run_elem is None:
            last_run = None
            last_idx = 0
            for idx, child in enumerate(children):
                if child.tag == f"{{{W}}}r":
                    last_run = child
                    last_idx = idx

            if last_run is None:
                target_run_elem = etree.Element(f"{{{W}}}r")
                etree.SubElement(target_run_elem, f"{{{W}}}t")
                p_elem.append(target_run_elem)
                run_position_in_parent = len(children)
                offset_in_run = 0
            else:
                target_run_elem = last_run
                t_elem = last_run.find(f"{{{W}}}t")
                offset_in_run = len(t_elem.text) if (t_elem is not None and t_elem.text) else 0
                run_position_in_parent = last_idx

        self._split_run_and_insert_ref(
            p_elem,
            target_run_elem,
            run_position_in_parent,
            offset_in_run,
            docx_fn_id,
        )

    def _split_run_and_insert_ref(self, p_elem, run_elem, run_pos: int, offset: int, docx_fn_id: int) -> None:
        """Split a run and insert a footnote-reference run between the halves."""
        t_elem = run_elem.find(f"{{{W}}}t")
        original_text = t_elem.text if (t_elem is not None and t_elem.text) else ""
        r_pr_elem = run_elem.find(f"{{{W}}}rPr")

        text_before = original_text[:offset]
        text_after = original_text[offset:]
        new_elements = []

        if text_before:
            r_before = etree.Element(f"{{{W}}}r")
            if r_pr_elem is not None:
                r_before.append(copy.deepcopy(r_pr_elem))
            t_before = etree.SubElement(r_before, f"{{{W}}}t")
            t_before.set(XML_SPACE, "preserve")
            t_before.text = text_before
            new_elements.append(r_before)

        r_ref = etree.Element(f"{{{W}}}r")
        r_pr_ref = etree.SubElement(r_ref, f"{{{W}}}rPr")
        if self.ref_style_id:
            r_style = etree.SubElement(r_pr_ref, f"{{{W}}}rStyle")
            r_style.set(f"{{{W}}}val", self.ref_style_id)
        else:
            vert_align = etree.SubElement(r_pr_ref, f"{{{W}}}vertAlign")
            vert_align.set(f"{{{W}}}val", "superscript")
        fn_ref = etree.SubElement(r_ref, f"{{{W}}}footnoteReference")
        fn_ref.set(f"{{{W}}}id", str(docx_fn_id))
        new_elements.append(r_ref)

        if text_after:
            r_after = etree.Element(f"{{{W}}}r")
            if r_pr_elem is not None:
                r_after.append(copy.deepcopy(r_pr_elem))
            t_after = etree.SubElement(r_after, f"{{{W}}}t")
            t_after.set(XML_SPACE, "preserve")
            t_after.text = text_after
            new_elements.append(r_after)

        p_elem.remove(run_elem)
        for index, element in enumerate(new_elements):
            p_elem.insert(run_pos + index, element)

    def insert_from_table(self, alignment_table: List[Dict]) -> Dict:
        """Insert footnotes using the alignment table and fallback location strategy."""
        success = []
        failed = []
        insertion_plan = []

        for entry in alignment_table:
            try:
                location = self._find_insertion_location(entry)
                if location:
                    insertion_plan.append(
                        {
                            "entry": entry,
                            "para_index": location["para_index"],
                            "char_pos": location["char_pos"],
                            "para": location["para"],
                        }
                    )
                else:
                    failed.append({"id": entry["footnote_id"], "reason": "location_not_found"})
            except Exception as exc:
                failed.append({"id": entry["footnote_id"], "reason": str(exc)})

        insertion_plan.sort(key=lambda item: (item["para_index"], item["char_pos"]), reverse=True)
        fn_tree, fn_part = self._get_or_create_footnotes_tree()
        next_id = self._get_next_footnote_id(fn_tree)

        for plan in insertion_plan:
            entry = plan["entry"]
            footnote_text = entry.get("chinese_footnote", "")

            while next_id in (2, 3):
                next_id += 1
            docx_fn_id = next_id
            next_id += 1

            try:
                self._add_footnote_to_tree(fn_tree, docx_fn_id, footnote_text)
                self._insert_footnote_at_location(plan["para"], plan["char_pos"], docx_fn_id)
                success.append(entry["footnote_id"])
            except Exception as exc:
                failed.append({"id": entry["footnote_id"], "reason": str(exc)})

        self._save_footnotes_tree(fn_tree, fn_part)
        return {"success": success, "failed": failed}

    def insert_from_markers(self, marker_map: Dict[str, str]) -> Dict:
        """
        Replace marker tokens like [[FN_3]] with real DOCX footnote references.

        This is used in the auto-translation flow where the Chinese draft is generated
        with inline marker tokens first.
        """
        success = []
        failed = []
        found_markers = set()

        fn_tree, fn_part = self._get_or_create_footnotes_tree()
        next_id = self._get_next_footnote_id(fn_tree)

        for para in self.doc.paragraphs:
            text = para.text
            matches = list(FOOTNOTE_MARKER_PATTERN.finditer(text))
            if not matches:
                continue

            for match in reversed(matches):
                marker_name = match.group(1)
                found_markers.add(marker_name)
                footnote_text = marker_map.get(marker_name)
                if not footnote_text:
                    failed.append({"id": marker_name, "reason": "marker_text_missing"})
                    continue

                para.text = para.text[:match.start()] + para.text[match.end():]

                while next_id in (2, 3):
                    next_id += 1
                docx_fn_id = next_id
                next_id += 1

                try:
                    self._add_footnote_to_tree(fn_tree, docx_fn_id, footnote_text)
                    self._insert_footnote_at_location(para, match.start(), docx_fn_id)
                    success.append(marker_name)
                except Exception as exc:
                    failed.append({"id": marker_name, "reason": str(exc)})

        for marker_name in marker_map:
            if marker_name not in found_markers:
                failed.append({"id": marker_name, "reason": "marker_not_found_in_document"})

        self._save_footnotes_tree(fn_tree, fn_part)
        return {"success": success, "failed": failed}

    def _find_insertion_location(self, entry: Dict) -> Optional[Dict]:
        """Locate the insertion point using the existing 5-level fallback strategy."""
        chinese_word = entry.get("chinese_word", "")
        context_before = entry.get("context_before", "")
        context_after = entry.get("context_after", "")
        chinese_sentence = entry.get("chinese_sentence", "")
        occurrence = entry.get("word_occurrence", 1)
        footnote_id = entry["footnote_id"]

        full_pattern = context_before + chinese_word + context_after
        partial_after = chinese_word + context_after

        if chinese_sentence:
            for index, para in enumerate(self.paragraphs):
                if chinese_sentence in para.text:
                    match_pos = self._find_nth_occurrence(para.text, full_pattern, 1)
                    if match_pos != -1:
                        char_pos = match_pos + len(context_before) + len(chinese_word)
                        return {"para_index": index, "char_pos": char_pos, "para": para}
                    break

        count = 0
        for index, para in enumerate(self.paragraphs):
            if full_pattern in para.text:
                count += 1
                if count == occurrence:
                    match_pos = para.text.find(full_pattern)
                    char_pos = match_pos + len(context_before) + len(chinese_word)
                    return {"para_index": index, "char_pos": char_pos, "para": para}

        if chinese_sentence and partial_after:
            for index, para in enumerate(self.paragraphs):
                if chinese_sentence in para.text and partial_after in para.text:
                    match_pos = para.text.find(partial_after)
                    char_pos = match_pos + len(chinese_word)
                    print(f"[FALLBACK-3] Footnote {footnote_id}: matched word+after")
                    return {"para_index": index, "char_pos": char_pos, "para": para}

        if chinese_sentence and chinese_word:
            for index, para in enumerate(self.paragraphs):
                if chinese_sentence in para.text:
                    match_pos = self._find_nth_occurrence(para.text, chinese_word, occurrence)
                    if match_pos != -1:
                        char_pos = match_pos + len(chinese_word)
                        print(f"[FALLBACK-4] Footnote {footnote_id}: matched by word only")
                        return {"para_index": index, "char_pos": char_pos, "para": para}
                    break

        if chinese_sentence and chinese_word:
            def bigram_overlap(left: str, right: str) -> float:
                left_bigrams = {left[i:i + 2] for i in range(len(left) - 1)} if len(left) > 1 else set()
                right_bigrams = {right[i:i + 2] for i in range(len(right) - 1)} if len(right) > 1 else set()
                if not left_bigrams or not right_bigrams:
                    return 0.0
                return len(left_bigrams & right_bigrams) / len(left_bigrams | right_bigrams)

            scores = [(bigram_overlap(chinese_sentence, para.text), index) for index, para in enumerate(self.paragraphs)]
            scores.sort(reverse=True)

            for score, index in scores[:3]:
                if score < 0.15:
                    break
                para = self.paragraphs[index]
                match_pos = self._find_nth_occurrence(para.text, chinese_word, occurrence)
                if match_pos != -1:
                    char_pos = match_pos + len(chinese_word)
                    print(f"[FALLBACK-5] Footnote {footnote_id}: fuzzy paragraph match (score={score:.2f})")
                    return {"para_index": index, "char_pos": char_pos, "para": para}

        print(f"[ERROR] Footnote {footnote_id}: all location strategies failed.")
        return None

    def _find_nth_occurrence(self, text: str, pattern: str, n: int) -> int:
        """Return the index of the nth occurrence or -1 if absent."""
        start = 0
        for _ in range(n):
            pos = text.find(pattern, start)
            if pos == -1:
                return -1
            start = pos + 1
        return text.find(pattern, start - 1)

    def save(self, output_path: str) -> None:
        """Save the modified document."""
        self.doc.save(output_path)
        print(f"[OK] Saved DOCX: {output_path}")
