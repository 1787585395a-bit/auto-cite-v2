"""
Helpers for generating simple DOCX drafts.
"""
from __future__ import annotations

from docx import Document


def build_docx_from_text(text: str, output_path: str) -> None:
    """Create a simple DOCX from plain text paragraphs."""
    doc = Document()

    paragraphs = [para.strip() for para in text.replace("\r\n", "\n").split("\n\n")]
    non_empty = [para for para in paragraphs if para]

    if not non_empty:
        doc.add_paragraph("")
    else:
        for para in non_empty:
            doc.add_paragraph(para)

    doc.save(output_path)
